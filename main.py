import argparse
import openai
import json
import time
import os
import pyfiglet
from dotenv import load_dotenv
from docx import Document

# Thiết lập các thông số cho ChatGPT
load_dotenv()
model_name = "gpt-3.5-turbo"

# Hàm để tạo các câu trả lời từ câu hỏi đầu vào
def generate_answers(questions):
    response = openai.ChatCompletion.create(
        model=model_name,
        messages=[
            {"role": "system", "content": "Bạn là một trợ lý hữu ích."},
            {"role": "user", "content": questions}
        ],
        temperature=0.7,
    )
    answer = response.choices[0].message.content.strip()

    return remove_first_line(answer)

def remove_first_line(answer):
    lines = answer.split('\n')

    first_line = lines[0]
    # Kiểm tra và loại bỏ dòng đầu tiên nếu chứa các từ "tóm tắt", "dưới đây" hoặc "sau đây"
    if any(word in first_line.lower() for word in ["tóm tắt", "dưới đây", "sau đây", ":"]):
        lines = lines[1:]

    result = '\n'.join(lines)  # Kết hợp lại các dòng còn lại thành một đoạn văn bản

    return result

# Đọc nội dung câu hỏi từ file .txt
def read_questions_from_file(input_path, output_path):
    try:
        file_extension_txt = os.path.splitext(input_path)[1]
        file_extension_docx = os.path.splitext(output_path)[1]

        if file_extension_txt == ".txt" and file_extension_docx == ".docx":
            with open(input_path, "r", encoding="utf-8") as file:
                questions = file.readlines()
                questions = [question.strip() for question in questions]
            return questions
        else:
            print("Error: Sai định dạng file input hoặc output")
            exit()
    except FileNotFoundError:
        print(f'Error: Không có file {input_path} trong thư mục')
        exit()

# Chạy chương trình với các câu hỏi từ file .txt
def run_chatbot_with_file(input_path, output_path):
    questions = read_questions_from_file(input_path, output_path)
    document = Document()

    for question in questions:
        print("Câu hỏi:", question)
        print("+ Lấy thông tin câu trả lời...")
        try:
            answer = generate_answers(question)
        except openai.error.RateLimitError as e:
            retry_time = e.retry_after if hasattr(e, 'retry_after') else 30
            print(f"Rate limit quá giới hạn. thử lại sau {retry_time} giây...")
            time.sleep(retry_time)
            answer = generate_answers(question)
        
        question_paragraph = document.add_paragraph()
        question_paragraph.add_run("Câu hỏi: ").bold = True
        question_paragraph.add_run(question)

        answer_paragraph = document.add_paragraph()
        answer_paragraph.add_run("Câu trả lời: ").bold = True
        answer_paragraph.add_run(answer)

        document.add_paragraph("")  # Thêm một dòng trống giữa các câu hỏi

    document.save(output_path)

def get_api_key(config):
    with open(config, "r") as file:
        json_data = json.load(file)

    return json_data['OPENAI_API_KEY_CREATE']

if __name__ == "__main__":
    ascii_banner = pyfiglet.figlet_format("ChatGPT Crawler")
    print(ascii_banner)
    print('                                               -- Created by Kieu Thanh Nam -- \n')
    parser = argparse.ArgumentParser()
    parser.add_argument("-c", "--config", type=str, help="Path to the file config", required=True)
    parser.add_argument("-f", "--file", type=str, help="Path to the .txt file containing questions", required=True)
    parser.add_argument("-o", "--output", type=str, help="Path to the output .docx file", required=True)
    args = parser.parse_args()

    try:
        openai.api_key = get_api_key(args.config)
        run_chatbot_with_file(args.file, args.output)
        print('Success: Crawl thông tin thành công!')
    except KeyboardInterrupt:
        print("Lỗi xảy ra khi cố thoát chương trình, vui lòng chạy lại")
        exit()
    except Exception as e:
        print(f"Có lỗi xảy ra: {e}")
